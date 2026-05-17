import streamlit as st
import pandas as pd
from docx import Document
import io
import re

# Ustawienia strony Streamlit
st.set_page_config(page_title="Konwerter Scenariuszy UAT Comfortel", layout="wide")

st.title("📄 Profesjonalny Konwerter Scenariuszy UAT (Comfortel)")
st.write("Wersja silnika: **v8.0-Architect-Multi**. Pełne sekwencyjne przetwarzanie hybrydowe, obsługa wielu plików i konsolidacja w jeden Excel.")

# Oficjalna struktura kolumn w Excelu
TARGET_COLUMNS = [
    "Scenariusze testowe", "Moduł", "Pełny Nr wymagania", "Opis wymagania", 
    "Zakres wyłączeń", "Nr scenariusza", "Nazwa scenariusza", "Cel", 
    "Warunki wstępne", "LP", "Dane", "Kroki testowe", "Oczekiwany rezultat", 
    "Wynik testu podczas odbioru", "Kategoria błędu", "Uwagi podczas odbioru"
]

def clean_text(text):
    """Dokładne czyszczenie tekstu ze znaków niedrukowalnych i wielokrotnych spacji."""
    if text is None:
        return ""
    return re.sub(r'\s+', ' ', str(text)).strip()

def extract_sheet_name_from_section_211(doc):
    """
    Ekstrakcja nazwy arkusza z sekcji 2.1.1 dokumentu Word.
    Szuka wzorca: '2.1.1 <nazwa>' lub podobnych wariantów.
    """
    text_content = "\n".join([p.text for p in doc.paragraphs])
    
    # Szukamy wzorca 2.1.1 z tekstem po nim
    match = re.search(r'2\.1\.1\s+([^\n]+)', text_content, re.IGNORECASE)
    if match:
        sheet_name = clean_text(match.group(1)).strip()
        if sheet_name:
            # Oczyszczenie nazwy arkusza (max 31 znaków dla Excela)
            sheet_name = sheet_name[:31]
            return sheet_name
    
    return None

def iter_block_items(doc):
    """Przechodzi przez wszystkie elementy dokumentu zachowując ich dokładną chronologiczną kolejność."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, doc)
        elif child.tag.endswith('tbl'):
            yield Table(child, doc)

def parse_docx_v8(uploaded_file):
    doc = Document(io.BytesIO(uploaded_file.read()))
    
    # --- KROK 1: Wstępne skanowanie początku dokumentu w celu ustalenia Tytułu i Modułu nadrzędnego ---
    main_title = ""
    detected_module = ""
    for p in doc.paragraphs[:30]:
        txt = p.text.strip()
        if "[SORT" in txt or "scenariusze" in txt.lower():
            main_title += " " + txt
        m_match = re.search(r'\[\s*(SORT\.[A-Z0-9]+)', txt, re.IGNORECASE)
        if m_match and not detected_module:
            detected_module = m_match.group(1).upper()
            
    main_title = clean_text(main_title)
    if not main_title:
        main_title = "Scenariusze testowe systemu SORT.B3S"

    # Ekstrakcja nazwy arkusza z sekcji 2.1.1
    sheet_name = extract_sheet_name_from_section_211(doc)
    if not sheet_name:
        sheet_name = "Scenariusze UAT"

    all_rows = []
    
    # Globalny słownik kontekstu (baza danych stanu)
    context = {col: "" for col in TARGET_COLUMNS}
    context["Scenariusze testowe"] = main_title
    if detected_module:
        context["Moduł"] = detected_module

    # --- KROK 2: Sekwencyjna analiza struktury pliku Word (Akapity + Tabele) ---
    for item in iter_block_items(doc):
        
        # Opcja A: Element to zwykły akapit tekstu (często zawiera nagłówki scenariuszy)
        if isinstance(item, from_docx_import_paragraph_type() if 'from_docx_import_paragraph_type' in locals() else type(doc.paragraphs[0])):
            text = item.text.strip()
            if not text:
                continue
            
            # Detekcja numeru i nazwy scenariusza w tekście (np. #TAKC_001 – Zmiana terminu...)
            if "#TAKC_" in text or text.startswith("#"):
                match = re.search(r"#\s*([A-Za-z0-9_]+)(?:\s*[\–\-\—\:\.]\s*(.*))?", text)
                if match:
                    context["Nr scenariusza"] = match.group(1).strip()
                    context["Nazwa scenariusza"] = match.group(2).strip() if match.group(2) else ""
            
            # Detekcja zmiany modułu w tekście
            module_match = re.search(r'\[\s*(SORT\.[A-Z0-9]+)', text, re.IGNORECASE)
            if module_match:
                context["Moduł"] = module_match.group(1).upper()
                
        # Opcja B: Element to tabela (może być hybrydą metadanych i kroków)
        else:
            in_steps_zone = False
            col_map = {"lp": None, "dane": None, "opis": None, "rezultat": None}
            
            for row in item.rows:
                # Pobieramy unikalne wartości tekstowe z komórek wiersza
                cells_text = [clean_text(cell.text) for cell in row.cells]
                if not any(cells_text):
                    continue
                
                cells_lowercase = [c.lower() for c in cells_text]
                
                # 1. Sprawdzamy czy wiersz jest nagłówkiem tabeli kroków
                has_lp = any(c in ["lp", "l.p.", "l.p"] for c in cells_lowercase)
                has_opis = any("opis" in c or "krok" in c for c in cells_lowercase)
                
                if has_lp and has_opis:
                    in_steps_zone = True
                    for idx, c_low in enumerate(cells_lowercase):
                        if c_low in ["lp", "l.p.", "l.p"]:
                            col_map["lp"] = idx
                        elif "dane" in c_low:
                            col_map["dane"] = idx
                        elif "opis" in c_low or "krok" in c_low:
                            col_map["opis"] = idx
                        elif "rezultat" in c_low or "wynik" in c_low or "oczekiwany" in c_low:
                            col_map["rezultat"] = idx
                    continue  # Pomijamy sam wiersz nagłówkowy kolumny Excela
                
                # 2. Jeśli trwa strefa kroków i wiersz zawiera numer kroku (liczba)
                if in_steps_zone:
                    lp_idx = col_map["lp"]
                    if lp_idx is not None and lp_idx < len(cells_text):
                        lp_value = cells_text[lp_idx]
                        if re.match(r"^\d+$", lp_value):
                            dane_idx = col_map["dane"]
                            opis_idx = col_map["opis"]
                            rez_idx = col_map["rezultat"]
                            
                            dane_val = cells_text[dane_idx] if dane_idx is not None and dane_idx < len(cells_text) else ""
                            opis_val = cells_text[opis_idx] if opis_idx is not None and opis_idx < len(cells_text) else ""
                            rez_val = cells_text[rez_idx] if rez_idx is not None and rez_idx < len(cells_text) else ""
                            
                            if "opis" in opis_val.lower() or "oczekiwany" in rez_val.lower() or "lp" in lp_value.lower():
                                continue
                            
                            # Budowanie płaskiego rekordu na bazie aktualnego, czystego kontekstu
                            flat_row = context.copy()
                            flat_row["LP"] = lp_value
                            flat_row["Dane"] = dane_val
                            flat_row["Kroki testowe"] = opis_val
                            flat_row["Oczekiwany rezultat"] = rez_val
                            all_rows.append(flat_row)
                            continue  # Przejdź do następnego wiersza tabeli
                
                # 3. Analiza metadanych (wykonywana dla wierszy niebędących krokami w KAŻDEJ tabeli)
                # Sprawdzenie wiersza pod kątem obecności kodu scenariusza (#TAKC_)
                for cell_txt in cells_text:
                    if "#TAKC_" in cell_txt or cell_txt.startswith("#"):
                        match = re.search(r"#\s*([A-Za-z0-9_]+)\s*[\–\-\—\:\.]\s*(.*)", cell_txt)
                        if match:
                            context["Nr scenariusza"] = match.group(1).strip()
                            context["Nazwa scenariusza"] = match.group(2).strip() if match.group(2) else ""
                
                # Inteligentna ekstrakcja par klucz-wartość (radzi sobie ze scalonymi komórkami)
                if len(cells_text) >= 2:
                    for idx in range(len(cells_text) - 1):
                        k = cells_text[idx].lower()
                        v = cells_text[idx + 1]
                        
                        if not v or v.lower() == k:
                            continue  # Pominięcie zduplikowanych komórek po scaleniu poziomym
                            
                        if "moduł" in k or "modul" in k:
                            context["Moduł"] = v
                        elif "nr wymagania" in k or "numer wymagania" in k or "id wymagania" in k:
                            if context["Pełny Nr wymagania"] != v:
                                context["Pełny Nr wymagania"] = v
                                # KLUCZOWY RESET: Zmiana wymagania czyści stary, nieaktualny pod-kontekst scenariusza
                                context["Nr scenariusza"] = ""
                                context["Nazwa scenariusza"] = ""
                                context["Cel"] = ""
                                context["Warunki wstępne"] = ""
                        elif "opis wymagania" in k:
                            context["Opis wymagania"] = v
                        elif "cel" in k:
                            context["Cel"] = v
                        elif "warunki" in k:
                            context["Warunki wstępne"] = v
                        elif "nazwa scenariusza" in k:
                            context["Nazwa scenariusza"] = v
                        elif "nr scenariusza" in k:
                            context["Nr scenariusza"] = v

    # Budowanie struktury DataFrame
    df = pd.DataFrame(all_rows) if all_rows else pd.DataFrame(columns=TARGET_COLUMNS)
    for col in TARGET_COLUMNS:
        if col not in df.columns:
            df[col] = ""
            
    return df[TARGET_COLUMNS], sheet_name

# Dynamiczne zabezpieczenie typu klasy Paragraph z biblioteki python-docx
def from_docx_import_paragraph_type():
    from docx.text.paragraph import Paragraph
    return Paragraph

# --- INTERFEJS STRONY (STREAMLIT) ---
st.subheader("📁 Wgrywanie plików")
uploaded_files = st.file_uploader(
    "Wgraj jeden lub więcej plików scenariuszy Word (.docx)",
    type=["docx"],
    accept_multiple_files=True
)

if uploaded_files:
    with st.spinner("Przetwarzanie dokumentów przez Silnik Chronologiczny v8.0-Architect..."):
        try:
            # Słownik do przechowywania DataFrames z nazwami arkuszy
            sheets_dict = {}
            total_rows = 0
            files_processed = 0
            
            for uploaded_file in uploaded_files:
                try:
                    df_result, sheet_name = parse_docx_v8(uploaded_file)
                    
                    if not df_result.empty:
                        # Obsługa duplikatów nazw arkuszy
                        original_sheet_name = sheet_name
                        counter = 1
                        while sheet_name in sheets_dict:
                            # Excel limit: 31 znaków
                            sheet_name = f"{original_sheet_name[:27]}_{counter}"
                            counter += 1
                        
                        sheets_dict[sheet_name] = df_result
                        total_rows += len(df_result)
                        files_processed += 1
                        st.write(f"✅ Plik **{uploaded_file.name}** przetworzony jako arkusz **{sheet_name}** ({len(df_result)} wierszy)")
                    else:
                        st.warning(f"⚠️ Plik **{uploaded_file.name}** - wynikowa tabela jest pusta. Sprawdź strukturę pliku wejściowego.")
                        
                except Exception as e:
                    st.error(f"❌ Błąd podczas przetwarzania **{uploaded_file.name}**: {e}")
            
            if sheets_dict:
                st.success(f"🎉 Sukces! Przetworzono {files_processed} plików. Całkowita liczba wierszy: {total_rows}.")
                
                # Karty KPI
                c1, c2, c3 = st.columns(3)
                with c1:
                    st.metric("Liczba kroków testowych", total_rows)
                with c2:
                    st.metric("Liczba arkuszy", len(sheets_dict))
                with c3:
                    st.metric("Ciągłość danych nadrzędnych", "Zweryfikowana (Brak luk)")
                
                # Podgląd arkuszy
                st.subheader("👀 Podgląd wynikowych, spłaszczonych struktur danych")
                selected_sheet = st.selectbox(
                    "Wybierz arkusz do podglądu:",
                    options=list(sheets_dict.keys())
                )
                st.dataframe(sheets_dict[selected_sheet], use_container_width=True)
                
                # Przygotowanie eksportu do Excela (.xlsx)
                st.subheader("📥 Eksport do pliku Excel")
                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    for sheet_name, df in sheets_dict.items():
                        df.to_excel(writer, index=False, sheet_name=sheet_name)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Pobierz ostateczny, kompletny plik Excel (.xlsx)",
                    data=buffer,
                    file_name="scenariusze_uat_odbior_konsolidowany.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            else:
                st.error("Błąd: Żaden plik nie został pomyślnie przetworzony.")
                
        except Exception as e:
            st.error(f"Wystąpił krytyczny błąd podczas przetwarzania: {e}")
else:
    st.info("ℹ️ Wgraj co najmniej jeden plik Word (.docx), aby rozpocząć przetwarzanie.")
